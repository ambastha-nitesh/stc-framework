"""Shared helpers for building STC Framework Word documents.

Matches the house style already used in the existing .docx files:
Title page, Table of Contents placeholder, numbered H1 sections, H2
sub-sections, body paragraphs, optional tables, optional call-out
blocks.

Each doc generator calls ``build(path, sections)`` with a structured
outline. The helper produces byte-for-byte reproducible output so
re-running the script is idempotent.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches


# --------------------------------------------------------------------------
# Style primitives
# --------------------------------------------------------------------------


def _shade_cell(cell, hex_fill: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def add_title_block(
    doc: Document,
    *,
    title: str,
    subtitle: str,
    tagline: str,
    classification: str = "INTERNAL",
    author: str = "Nitesh Ambastha",
    version: str = "Version 2.0 \u2014 April 2026",
    doc_id: str | None = None,
) -> None:
    """The cover-page band that all STC docs share."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("STC FRAMEWORK")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x63, 0x66, 0xF1)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.italic = True
    r.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(tagline)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

    doc.add_paragraph()

    if classification:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"CLASSIFICATION: {classification}")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

    meta_lines = [f"Author: {author}", version]
    if doc_id:
        meta_lines.insert(0, f"Document ID: {doc_id}")
    for line in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    doc.add_page_break()


def add_toc(doc: Document, entries: list[str]) -> None:
    doc.add_heading("Table of Contents", level=1)
    for i, title in enumerate(entries, start=1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. {title}")
        r.font.size = Pt(11)
    doc.add_page_break()


def add_callout(doc: Document, label: str, body: str, *, color: str = "DBEAFE") -> None:
    """A shaded single-row 1x1 table used as a callout."""
    t = doc.add_table(rows=1, cols=1)
    t.autofit = True
    cell = t.rows[0].cells[0]
    _shade_cell(cell, color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    first = cell.paragraphs[0]
    r = first.add_run(f"{label}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

    p = cell.add_paragraph()
    r = p.add_run(body)
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    doc.add_paragraph()


def add_kv_table(doc: Document, rows: list[tuple[str, str]], *, key_width: float = 2.0) -> None:
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(rows):
        t.rows[i].cells[0].text = k
        t.rows[i].cells[1].text = v
        for p in t.rows[i].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
    # widths
    try:
        for row in t.rows:
            row.cells[0].width = Inches(key_width)
    except Exception:
        pass
    doc.add_paragraph()


def add_grid_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    # header
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
        _shade_cell(hdr[i], "E0E7FF")
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = val
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_para(doc: Document, text: str, *, italic: bool = False, bold: bool = False) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    if italic:
        r.italic = True
    if bold:
        r.bold = True


def add_code(doc: Document, code: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


# --------------------------------------------------------------------------
# Structured doc schema
# --------------------------------------------------------------------------


@dataclass
class Section:
    """A single top-level section with optional sub-sections and content."""

    title: str
    body: list[dict[str, Any]] = field(default_factory=list)


def render_section(doc: Document, section: Section, *, level: int = 1) -> None:
    doc.add_heading(section.title, level=level)
    for blk in section.body:
        t = blk.get("type")
        if t == "para":
            add_para(doc, blk["text"], italic=blk.get("italic", False), bold=blk.get("bold", False))
        elif t == "bullets":
            add_bullets(doc, blk["items"])
        elif t == "numbered":
            add_numbered(doc, blk["items"])
        elif t == "kv":
            add_kv_table(doc, blk["rows"], key_width=blk.get("key_width", 2.0))
        elif t == "grid":
            add_grid_table(doc, blk["headers"], blk["rows"])
        elif t == "callout":
            add_callout(doc, blk["label"], blk["text"], color=blk.get("color", "DBEAFE"))
        elif t == "code":
            add_code(doc, blk["code"])
        elif t == "sub":
            render_section(doc, Section(title=blk["title"], body=blk["body"]), level=level + 1)
        elif t == "pagebreak":
            doc.add_page_break()
        else:
            raise ValueError(f"Unknown block type: {t!r}")


def build(
    path: str | Path,
    *,
    title: str,
    subtitle: str,
    tagline: str,
    classification: str = "INTERNAL",
    version: str = "Version 2.0 \u2014 April 2026",
    doc_id: str | None = None,
    author: str = "Nitesh Ambastha",
    toc: list[str],
    sections: list[Section],
) -> None:
    doc = Document()
    # Narrow margins so tables fit.
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    # Base body style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_title_block(
        doc,
        title=title,
        subtitle=subtitle,
        tagline=tagline,
        classification=classification,
        version=version,
        doc_id=doc_id,
        author=author,
    )
    add_toc(doc, toc)
    for section in sections:
        render_section(doc, section)
    doc.save(str(path))


# --------------------------------------------------------------------------
# Convenience: common blocks for STC docs
# --------------------------------------------------------------------------


def session_changes_callout() -> dict[str, Any]:
    return {
        "type": "callout",
        "label": "April 2026 update \u2014 regulated-environment hardening",
        "text": (
            "This document has been updated to reflect the Round-1 and "
            "Round-2 staff-review remediations: HMAC-SHA256 audit chain, "
            "WORM-compatible audit backend, ed25519-signed declarative "
            "spec, strict production startup invariants, per-event-class "
            "retention policies, clock-safe per-tenant budget tracker, "
            "idempotency-aware erasure, and 255 regression tests across "
            "six audit suites (test_security, test_privacy, "
            "test_observability, test_enterprise, test_staff_review, "
            "test_staff_review_round2)."
        ),
        "color": "FEF3C7",
    }
